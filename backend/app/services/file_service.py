"""File processing service for extracting text from uploaded files."""

import logging
import base64
from io import BytesIO
from typing import Optional

logger = logging.getLogger(__name__)


class FileService:
    """Service for processing and extracting text from various file types."""
    
    @staticmethod
    def extract_text_from_file(file_content: bytes, filename: str, content_type: str) -> Optional[str]:
        """
        Extract text content from various file types.
        
        Supports: .txt, .pdf, .docx, .doc, .xlsx
        
        Args:
            file_content: Raw file bytes
            filename: Name of the file (used for extension detection)
            content_type: MIME type of the file
            
        Returns:
            Extracted text content, or None if extraction failed
        """
        try:
            if not file_content:
                logger.warning(f"[FileService] Empty file content for {filename}")
                return "[Empty file - no content to extract]"
            
            # Get file extension
            ext = filename.lower().split('.')[-1] if '.' in filename else ''
            logger.info(f"[FileService] Processing {filename}, extension: {ext}, size: {len(file_content)} bytes")
            logger.info(f"[FileService] Content-Type header: '{content_type}'")
            
            # Handle Image files (jpg, png, gif, webp) - CHECK THIS FIRST
            if ext in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                logger.info(f"[FileService] ✅ MATCHED by extension - Processing as IMAGE file")
                result = FileService._process_image_for_analysis(file_content, filename, ext, content_type)
                logger.info(f"[FileService] Image processing returned: {len(str(result)) if result else 0} chars")
                return result
            elif 'image' in content_type.lower():
                logger.info(f"[FileService] ✅ MATCHED by content-type - Processing as IMAGE file")
                result = FileService._process_image_for_analysis(file_content, filename, ext, content_type)
                logger.info(f"[FileService] Image processing returned: {len(str(result)) if result else 0} chars")
                return result
            
            # Handle text files
            elif ext == 'txt' or 'text' in content_type.lower():
                logger.info(f"[FileService] Processing as TXT file")
                return FileService._extract_text_from_txt(file_content)
            
            # Handle PDF files
            elif ext == 'pdf' or 'pdf' in content_type.lower():
                logger.info(f"[FileService] Processing as PDF file")
                return FileService._extract_text_from_pdf(file_content)
            
            # Handle Word documents (.docx)
            elif ext == 'docx' or 'openxmlformats-officedocument.wordprocessingml' in content_type.lower():
                logger.info(f"[FileService] Processing as DOCX file")
                return FileService._extract_text_from_docx(file_content)
            
            # Handle older Word documents (.doc)
            elif ext == 'doc' or 'msword' in content_type.lower():
                logger.info(f"[FileService] Processing as DOC file")
                return FileService._extract_text_from_doc(file_content)
            
            # Handle Excel files
            elif ext in ['xlsx', 'xls'] or 'spreadsheet' in content_type.lower():
                logger.info(f"[FileService] Processing as Excel file")
                return FileService._extract_text_from_excel(file_content)
            
            # Handle Video files (mp4, webm, mov, avi, mkv)
            elif ext in ['mp4', 'webm', 'mov', 'avi', 'mkv', 'flv'] or 'video' in content_type.lower():
                logger.info(f"[FileService] Processing as VIDEO file")
                return FileService._process_video_file(filename, ext, content_type)
            
            # Handle Code files (python, javascript, etc.)
            elif ext in ['py', 'js', 'ts', 'tsx', 'jsx', 'cpp', 'c', 'java', 'cs', 'rb', 'go', 'rs', 'php', 'sql', 'html', 'css', 'json', 'xml', 'yaml', 'yml', 'md']:
                logger.info(f"[FileService] Processing as CODE file")
                return FileService._extract_code_file(file_content, filename, ext)
            
            # Handle CSV files (table data)
            elif ext == 'csv' or 'csv' in content_type.lower():
                logger.info(f"[FileService] Processing as CSV file")
                return FileService._extract_csv_data(file_content, filename)
            
            else:
                logger.warning(f"[FileService] Unsupported file type: {ext} ({content_type})")
                return f"[Unsupported file type: {ext}. Supported types: txt, pdf, docx, doc, xlsx, xls, jpg, png, gif, webp, mp4, webm, mov, avi, mkv, py, js, ts, csv, json, xml, md, sql, html, css]"
                
        except Exception as e:
            logger.error(f"[FileService] Error extracting text from {filename}: {str(e)}", exc_info=True)
            return f"[Error reading file {filename}: {str(e)}]"
    
    @staticmethod
    def _extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from .txt files."""
        try:
            text = file_content.decode('utf-8', errors='ignore').strip()
            logger.info(f"[FileService] Extracted {len(text)} characters from TXT file")
            return text if text else "[Empty text file]"
        except Exception as e:
            logger.error(f"[FileService] Error extracting from txt: {str(e)}")
            return "[Error reading text file]"
    
    @staticmethod
    def _extract_text_from_pdf(file_content: bytes) -> Optional[str]:
        """Extract text from PDF files."""
        try:
            import PyPDF2
            logger.info(f"[FileService] PyPDF2 imported successfully")
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            num_pages = len(pdf_reader.pages)
            logger.info(f"[FileService] PDF has {num_pages} pages")
            
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    text = page.extract_text()
                    if text and text.strip():
                        text_content.append(f"[Page {page_num + 1}]\n{text}")
                        logger.debug(f"[FileService] Extracted {len(text)} chars from page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"[FileService] Error extracting page {page_num}: {str(e)}")
            
            result = "\n\n".join(text_content) if text_content else "[No text found in PDF]"
            logger.info(f"[FileService] Extracted {len(result)} total characters from PDF")
            return result
            
        except ImportError:
            logger.error(f"[FileService] PyPDF2 not installed")
            return "[PDF support not available - PyPDF2 not installed. Please install: pip install PyPDF2]"
        except Exception as e:
            logger.error(f"[FileService] Error extracting from PDF: {str(e)}", exc_info=True)
            return f"[Error reading PDF file: {str(e)}]"
    
    @staticmethod
    def _extract_text_from_docx(file_content: bytes) -> Optional[str]:
        """Extract text from .docx files."""
        try:
            from docx import Document
            logger.info(f"[FileService] python-docx imported successfully")
            
            doc = Document(BytesIO(file_content))
            text_content = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(
                        cell.text.strip() for cell in row.cells
                    )
                    if row_text.strip():
                        text_content.append(row_text)
            
            result = "\n".join(text_content) if text_content else "[Empty Word document]"
            logger.info(f"[FileService] Extracted {len(result)} characters from DOCX")
            return result
            
        except ImportError:
            logger.error(f"[FileService] python-docx not installed")
            return "[DOCX support not available - python-docx not installed. Please install: pip install python-docx]"
        except Exception as e:
            logger.error(f"[FileService] Error extracting from docx: {str(e)}", exc_info=True)
            return f"[Error reading DOCX file: {str(e)}]"
    
    @staticmethod
    def _extract_text_from_doc(file_content: bytes) -> Optional[str]:
        """Extract text from old .doc files (Word 97-2003)."""
        try:
            # Try python-docx first (works for some .doc files)
            try:
                from docx import Document
                doc = Document(BytesIO(file_content))
                text_content = []
                
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
                
                result = "\n".join(text_content) if text_content else "[Empty Word document]"
                logger.info(f"[FileService] Successfully extracted from DOC using python-docx")
                return result
            except:
                # Fall back to docx2txt
                logger.info(f"[FileService] python-docx failed, trying docx2txt")
                import docx2txt
                import tempfile
                import os
                
                # docx2txt needs a file path, so write to temp file
                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as tmp:
                    tmp.write(file_content)
                    tmp_path = tmp.name
                
                try:
                    text = docx2txt.process(tmp_path)
                    result = text.strip() if text else "[Empty Word document]"
                    logger.info(f"[FileService] Extracted from DOC using docx2txt")
                    return result
                finally:
                    os.unlink(tmp_path)
            
        except ImportError as e:
            logger.error(f"[FileService] Required library not installed: {str(e)}")
            return "[DOC support not available - please install: pip install python-docx docx2txt]"
        except Exception as e:
            logger.error(f"[FileService] Error extracting from doc: {str(e)}", exc_info=True)
            return f"[Error reading DOC file: {str(e)}]"
    
    @staticmethod
    def _extract_text_from_excel(file_content: bytes) -> Optional[str]:
        """Extract text from Excel files (.xlsx, .xls)."""
        try:
            import openpyxl
            logger.info(f"[FileService] openpyxl imported successfully")
            
            workbook = openpyxl.load_workbook(BytesIO(file_content))
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                worksheet = workbook[sheet_name]
                text_content.append(f"[Sheet: {sheet_name}]")
                
                for row in worksheet.iter_rows(values_only=True):
                    row_text = ' | '.join(
                        str(cell) if cell is not None else '' 
                        for cell in row
                    ).strip()
                    if row_text:
                        text_content.append(row_text)
                
                text_content.append("")  # Separator between sheets
            
            result = "\n".join(text_content) if text_content else "[Empty Excel file]"
            logger.info(f"[FileService] Extracted {len(result)} characters from Excel")
            return result
            
        except ImportError:
            logger.error(f"[FileService] openpyxl not installed")
            return "[Excel support not available - openpyxl not installed. Please install: pip install openpyxl]"
        except Exception as e:
            logger.error(f"[FileService] Error extracting from Excel: {str(e)}", exc_info=True)
            return f"[Error reading Excel file: {str(e)}]"
    
    @staticmethod
    def _process_image_for_analysis(file_content: bytes, filename: str, ext: str, content_type: str) -> str:
        """
        Process image file for AI analysis by converting to base64.
        
        The image is encoded as base64 and formatted so the LLM can analyze it.
        
        Args:
            file_content: Raw image bytes
            filename: Name of the image file
            ext: File extension (jpg, png, etc.)
            content_type: MIME type of the image
            
        Returns:
            Formatted string with base64-encoded image data and instruction for LLM analysis
        """
        try:
            logger.info(f"[FileService] 🖼️ Processing image: {filename}")
            logger.info(f"[FileService] 🖼️ Image type: {content_type}, Size: {len(file_content)} bytes")
            
            # Convert to base64
            image_base64 = base64.b64encode(file_content).decode('utf-8')
            logger.info(f"[FileService] 🖼️ Image encoded to base64: {len(image_base64)} characters")
            
            # Determine MIME type
            mime_type = content_type if content_type and 'image' in content_type else f"image/{ext}"
            
            # Format for LLM: include image data URL format
            image_data_url = f"data:{mime_type};base64,{image_base64}"
            
            # Create a prompt instruction for the LLM to analyze the image
            analysis_instruction = f"""[IMAGE ANALYSIS REQUEST]
User has uploaded an image file: {filename}
Please analyze and describe the contents of this image in detail.
Image Format: {ext.upper()}

Image Data (base64):
{image_data_url}

Provide a detailed description including:
1. What you see in the image
2. Colors, lighting, and mood
3. Any text visible
4. Objects, people, or animals present
5. Overall composition and scene
[END IMAGE ANALYSIS REQUEST]"""
            
            logger.info(f"[FileService] 🖼️ ✅ Image processed successfully for analysis")
            logger.info(f"[FileService] 🖼️ Analysis instruction length: {len(analysis_instruction)} chars")
            logger.info(f"[FileService] 🖼️ Marker check - contains [IMAGE ANALYSIS REQUEST]: {'[IMAGE ANALYSIS REQUEST]' in analysis_instruction}")
            
            return analysis_instruction
            
        except Exception as e:
            logger.error(f"[FileService] 🖼️ Error processing image {filename}: {str(e)}", exc_info=True)
            return f"[Error processing image {filename}: {str(e)}]"
    
    @staticmethod
    def _process_video_file(filename: str, ext: str, content_type: str) -> str:
        """
        Process video file by providing analysis instructions to the LLM.
        
        Since we can't directly process videos, we provide metadata and ask
        the user for more details or suggest they describe key scenes.
        
        Args:
            filename: Name of the video file
            ext: File extension (mp4, webm, mov, etc.)
            content_type: MIME type of the video
            
        Returns:
            Formatted string with video metadata for LLM analysis
        """
        try:
            logger.info(f"[FileService] Processing video: {filename}")
            
            video_info = f"""[VIDEO FILE UPLOADED]
Filename: {filename}
Format: {ext.upper()}
MIME Type: {content_type}

Video Support Information:
- The video file has been uploaded successfully
- You can ask me to analyze the video content if you provide:
  1. Key scenes or timestamps to focus on
  2. Questions about specific parts of the video
  3. A description of what's shown in the video
  4. Subtitles or transcripts (if available)

You can also:
- Ask me to help extract key insights from the video
- Request a summary of the video content
- Ask questions about specific scenes or segments
[END VIDEO FILE INFORMATION]"""
            
            logger.info(f"[FileService] ✅ Video file processed successfully")
            return video_info
            
        except Exception as e:
            logger.error(f"[FileService] Error processing video {filename}: {str(e)}", exc_info=True)
            return f"[Error processing video {filename}: {str(e)}]"
    
    @staticmethod
    def _extract_code_file(file_content: bytes, filename: str, ext: str) -> str:
        """
        Extract and format code file for analysis.
        
        Args:
            file_content: Raw file bytes
            filename: Name of the code file
            ext: File extension (py, js, etc.)
            
        Returns:
            Formatted code block with syntax highlighting marker
        """
        try:
            logger.info(f"[FileService] Extracting code from {filename}")
            
            code_text = file_content.decode('utf-8', errors='ignore').strip()
            logger.info(f"[FileService] Extracted {len(code_text)} characters from code file")
            
            if not code_text:
                return f"[Empty code file: {filename}]"
            
            # Create markdown code block with language specified
            code_block = f"""[CODE FILE: {filename}]

```{ext}
{code_text}
```

[END CODE FILE]"""
            
            logger.info(f"[FileService] ✅ Code file processed successfully")
            return code_block
            
        except Exception as e:
            logger.error(f"[FileService] Error extracting code from {filename}: {str(e)}", exc_info=True)
            return f"[Error reading code file {filename}: {str(e)}]"
    
    @staticmethod
    def _extract_csv_data(file_content: bytes, filename: str) -> str:
        """
        Extract and format CSV file as a table for analysis.
        
        Args:
            file_content: Raw file bytes
            filename: Name of the CSV file
            
        Returns:
            Formatted table data (markdown format)
        """
        try:
            import csv
            logger.info(f"[FileService] Extracting CSV data from {filename}")
            
            csv_text = file_content.decode('utf-8', errors='ignore')
            reader = csv.reader(csv_text.splitlines())
            rows = list(reader)
            
            if not rows:
                return f"[Empty CSV file: {filename}]"
            
            # Create markdown table format
            headers = rows[0]
            table_lines = [
                f"[TABLE: {filename}]\n",
                "| " + " | ".join(headers) + " |",
                "| " + " | ".join(["---"] * len(headers)) + " |"
            ]
            
            # Add data rows (limit to first 100 rows to avoid overwhelming the LLM)
            for row in rows[1:101]:
                if len(row) == len(headers):
                    table_lines.append("| " + " | ".join(row) + " |")
            
            if len(rows) > 101:
                table_lines.append(f"\n[... and {len(rows) - 101} more rows]")
            
            table_lines.append(f"\n[END TABLE]")
            
            table_text = "\n".join(table_lines)
            logger.info(f"[FileService] ✅ CSV processed: {len(rows)} rows, {len(headers)} columns")
            return table_text
            
        except Exception as e:
            logger.error(f"[FileService] Error extracting CSV from {filename}: {str(e)}", exc_info=True)
            return f"[Error reading CSV file {filename}: {str(e)}]"
